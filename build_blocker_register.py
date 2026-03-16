"""Generate GMA Blocker Register as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

doc = Document()

# Page setup - landscape
for section in doc.sections:
    section.orientation = 1  # landscape
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

# Title
title = doc.add_heading('GMA — Blocker Register', level=0)
title.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

p = doc.add_paragraph()
p.add_run('Les Grands Moulins d\'Abidjan | Ivory Coast | Cash Management Transformation').bold = True
doc.add_paragraph(f'Prepared: {datetime.date.today().strftime("%B %d, %Y")} | Source: GMA Current State PPT (59 slides) + Transaction Data Analysis')
doc.add_paragraph()

# Severity color helper
SEV_COLORS = {
    'CRITICAL': RGBColor(0xc5, 0x30, 0x30),
    'HIGH': RGBColor(0xe5, 0x3e, 0x3e),
    'MEDIUM': RGBColor(0xd6, 0x9e, 0x2e),
    'LOW': RGBColor(0xa0, 0xae, 0xc0),
}

CAT_COLORS = {
    'Regulatory': RGBColor(0x2b, 0x6c, 0xb0),
    'Banking': RGBColor(0x27, 0x67, 0x49),
    'Operational': RGBColor(0xb7, 0x79, 0x1f),
    'Technology': RGBColor(0x6b, 0x46, 0xc1),
    'Legal / Tax': RGBColor(0x97, 0x4b, 0x00),
}

STATUS_MAP = {
    'Open': RGBColor(0xc5, 0x30, 0x30),
    'In Progress': RGBColor(0xd6, 0x9e, 0x2e),
    'Monitoring': RGBColor(0x2b, 0x6c, 0xb0),
}

blockers = [
    # (ID, Description, Category, Severity, Impact, Mitigation, Owner, Status, Source)
    ('BLK-01',
     'BCEAO 50% Prepayment Cap',
     'Regulatory',
     'CRITICAL',
     'Advance payments for wheat imports cannot exceed 50% of invoice value (BCEAO Instruction 02.07.2025 RFE, Art. 6(f)). Maximum prepayment horizon is 6 months with 3-month grace period. Full compliance mandatory since January 2026. Directly limits the velocity of cash repatriation to Eurafrique/SOG via commodity prepayments — GMA\'s lowest-cost repatriation method (0.71%).',
     'Maximize prepayment volumes within the 50% cap. Diversify repatriation via FX capex prepayments (also 0.71%) and intercompany loans (2–3% friction). Ensure all documentation (DPI, invoices, contracts) is prepared in advance to avoid batch delays.',
     'Cyril Edorh / Corp Treasury',
     'Open',
     'PPT Slide 22'),

    ('BLK-02',
     'FX Batch Rejection Risk at BCEAO',
     'Regulatory',
     'CRITICAL',
     'If any single payment in a bank\'s submitted FX batch is incorrect, the entire batch is rejected by BCEAO. When central bank liquidity is tight, BCEAO applies stricter scrutiny and finds more reasons to reject batches. This can delay all FX payments by days.',
     'Implement pre-submission validation checklist for all FX payment documentation. Separate high-value/time-sensitive payments into dedicated batches. Maintain relationships with multiple banks to diversify batch submission risk. Track BCEAO rejection rates by bank.',
     'Cyril Edorh',
     'Open',
     'PPT Slide 22 (Additional Comments)'),

    ('BLK-03',
     'Règlement 06/2024 — New FX Controls',
     'Regulatory',
     'HIGH',
     'New WAEMU foreign exchange regulation (effective Dec 20, 2024) replaces 2010 rules. Mandatory domiciliation of large trade transactions, repatriation of export proceeds, prior authorization for large outbound investments. 12 implementing Instructions issued July 2025. Impact on GMA\'s wheat import payments and intercompany transfers not fully assessed.',
     'Obtain and analyze all 12 implementing Instructions. Map each requirement to GMA\'s current payment flows. Identify any processes that need modification. Consult with local banks on their interpretation and compliance procedures.',
     'Corp Treasury / Local Counsel',
     'Open',
     'PPT Slides 45–46, Slide 8'),

    ('BLK-04',
     'FCY Account Opening Requires MoF + BCEAO Approval',
     'Regulatory',
     'MEDIUM',
     'Opening foreign currency accounts is subject to explicit approval from both the Ministry of Finance and BCEAO. This could delay any future account restructuring that requires FCY accounts (e.g., EUR or USD accounts for direct settlement).',
     'Assess whether FCY accounts are needed in target state. If yes, begin approval process early as it can take months. Document business justification per BCEAO requirements.',
     'Corp Treasury',
     'Monitoring',
     'PPT Slide 44'),

    ('BLK-05',
     'No PSD2-Equivalent — No Mandated Transfer Deadlines',
     'Regulatory',
     'MEDIUM',
     'WAEMU has no regulation mandating customer-level transfer execution deadlines (unlike EU PSD2). Banks are not legally required to credit customer accounts within any specific timeframe after interbank settlement. Customer credit timing depends entirely on individual bank processing and contractual SLAs.',
     'Negotiate explicit SLAs with each bank for crediting timelines. Use STAR-UEMOA for time-sensitive payments (immediate settlement). Leverage PI-SPI for instant payments where available. Document actual observed crediting times per bank.',
     'Corp Treasury',
     'Monitoring',
     'PPT Slide 15, Slide 53'),

    ('BLK-06',
     'AFG Bank — No Online Banking Platform',
     'Banking',
     'HIGH',
     'AFG (which acquired Banque Populaire) does not have an online banking platform. It is currently being developed for their customers. This prevents electronic payment initiation, automated statement retrieval, and any file-based integration. AFG is retained for remote depot presence (connected to post offices).',
     'Assess timeline for AFG online platform launch. In interim, maintain manual processes for AFG account. Evaluate whether remote collection can be migrated to PI-SPI or mobile money (Wave/Orange) as alternative to AFG physical presence.',
     'Cyril Edorh',
     'Open',
     'PPT Slide 28'),

    ('BLK-07',
     'Citibank — Highest Cost Bank Being Elevated as Primary',
     'Banking',
     'HIGH',
     'CITI is the most expensive of all GMA banks, yet is being positioned as the primary account for excess cash sweeps and CIO product. CITI has not honored prepayment pricing when liquidity is tight. Transitioning flows to CITI requires maintaining wheat prepayment commitments to other banks to preserve their pricing.',
     'Negotiate volume-based pricing with CITI tied to the CIO and sweep arrangement. Document pricing commitments to other banks for wheat flows. Conduct cost-benefit analysis: CIO savings vs higher per-transaction costs. Establish fallback if CITI pricing becomes uncompetitive.',
     'Corp Treasury',
     'In Progress',
     'PPT Slide 32'),

    ('BLK-08',
     'Saisie — Court-Ordered Account Freezes',
     'Banking',
     'HIGH',
     'Courts can order banks to freeze GMA accounts or transactions (Saisie) when vendors bring disputes. Frozen amount scope unclear (disputed amount vs entire account balance). Has occurred at most banks except BOA. Can immobilize working capital without warning.',
     'Clarify with legal counsel: is the freeze limited to the disputed amount or the entire account? Maintain sufficient liquidity across multiple banks to ensure operations continue if one account is frozen. Consider keeping higher balances at BOA (which has never been frozen). Document all vendor disputes proactively.',
     'Cyril Edorh / Legal',
     'Open',
     'PPT Slide 39, Slide 59'),

    ('BLK-09',
     'MTN Mobile Money Account Inactive',
     'Banking',
     'MEDIUM',
     'MTN Mobile Money account (XOF 6151) stopped reporting in November 2024. MTN has only ~20% mobile money market share in Ivory Coast. Account is functionally dormant and needs to be closed and replaced.',
     'Close MTN account. Evaluate Wave (~50% market share, American company) and Orange Money (market leader) as replacements. Cyril conducting audit. Consider PI-SPI integration as the long-term mobile collection channel instead of bilateral EMI relationship.',
     'Cyril Edorh',
     'In Progress',
     'PPT Slide 33'),

    ('BLK-10',
     'Trapped BOA Bank Guarantee (LG12/795)',
     'Banking',
     'LOW',
     'Bank guarantee LG12/795 issued by BOA for Port of Abidjan land (XOF 78.9M / $140K) expired in 2015 but BOA requires a release letter from the Port to remove it from their books. Port only provided a letter saying they don\'t recognize the guarantee — insufficient for the bank. Actions underway but unresolved for years.',
     'Escalate via legal counsel to obtain formal release letter from Port authority. If release cannot be obtained, explore whether BOA can write off the guarantee with the Port\'s non-recognition letter plus a legal opinion.',
     'Cyril Edorh / Legal',
     'In Progress',
     'PPT Slide 48'),

    ('BLK-11',
     'Cash-Intensive Depot Collection Model',
     'Operational',
     'CRITICAL',
     '17+ depots collect cash and paper-based instruments (checks, bills of exchange). Each depot has a cash threshold — when reached, cash must be physically transported to the nearest bank. This is GMA\'s #1 pain point (Slide 52). Creates security risk, float delay, reconciliation complexity, and limits to how quickly GMA can deploy collected cash.',
     'Pilot PI-SPI QR code / alias-based payment acceptance at high-volume depots. Evaluate Wave/Orange Money integration for remote depots. Target: reduce cash and paper collection by 50%+ over 24 months. Requires customer education and incentive program.',
     'Cyril Edorh / Commercial Team',
     'Open',
     'PPT Slides 39, 52'),

    ('BLK-12',
     'Check Processing Inefficiency',
     'Operational',
     'MEDIUM',
     'Future-dated checks stored in safe, manually reviewed daily by Joelle, endorsed and taken to bank on maturity date. Each bank visit takes at least 2 hours. Banks sometimes reject valid checks to collect unpaid-check fees. Declining volume but still operational burden.',
     'Track check volume trend (declining). Encourage customers to migrate to electronic payment methods (wire, PI-SPI). Negotiate with banks on check rejection practices. Consider consolidating check deposits to fewer banks to reduce visit frequency.',
     'Joelle / Cyril Edorh',
     'Monitoring',
     'PPT Slide 41'),

    ('BLK-13',
     'PI-SPI Adoption Still Early Stage',
     'Technology',
     'HIGH',
     'Only 15 Ivorian institutions on PI-SPI as of Feb 2026. Only ~221,000 customers registered a PI-SPI alias (Jan 2026). SocGen not yet on platform. PI-SPI is a single centralized hub — BCEAO outage cascades across all participants. GMA cannot fully rely on PI-SPI for collection modernization or payroll consolidation until adoption matures.',
     'Monitor PI-SPI participant list and adoption metrics. Begin with use cases where counterparties are already on PI-SPI (SIB, BOA). Maintain fallback payment channels. Pilot payroll consolidation with SIB/BOA employees first. Track SocGen onboarding timeline.',
     'Corp Treasury',
     'Monitoring',
     'PPT Slides 40, 54, Slide 13 (Risk Register)'),

    ('BLK-14',
     'No IT2 / IREC Readiness Assessment',
     'Technology',
     'CRITICAL',
     'ION IT2 TMS and BofA IREC bank connectivity have not been assessed for GMA. No configuration plan, no file format mapping, no integration architecture defined. This is a prerequisite for D.05 (Target Design) and D.07 (Technology Strategy). Without it, implementation planning (D.08) cannot begin.',
     'Conduct IT2 readiness assessment: entity structure, account mapping, deal types, payment workflows, accounting integration. Assess IREC availability in Ivory Coast. Define bank connectivity strategy per bank (SWIFT, H2H, API, manual). Map current bank file formats.',
     'Corp Treasury / IT',
     'Open',
     'Gap Analysis D.04, D.07'),

    ('BLK-15',
     'Dividend Repatriation — 15% WHT and Annual Window',
     'Legal / Tax',
     'MEDIUM',
     'Dividend payment is the most expensive repatriation method (15% permanent withholding tax). GMA has only one opportunity per year to declare a dividend, with strict deadlines (must complete by September 30). This limits repatriation flexibility and creates a high tax cost for cash extraction.',
     'Maximize lower-cost repatriation methods (commodity prepayments, FX capex) before resorting to dividends. Evaluate intercompany loan structures for interim liquidity. Investigate whether treaty relief applies (France-CI tax treaty). Ensure dividend declaration timeline is tracked and met annually.',
     'Corp Treasury / Tax',
     'Monitoring',
     'PPT Slide 22'),

    ('BLK-16',
     'Intercompany Loan Transfer Pricing Risk',
     'Legal / Tax',
     'MEDIUM',
     'Intercompany loans between WAEMU entities carry 2–3% annual tax friction (non-deductible tax on interest income). If outstanding too long, transfer pricing concerns arise under OHADA rules. Requires MoF declaration, registration fee, and board authorization.',
     'Establish clear loan terms with arm\'s-length interest rates. Set maximum duration limits. Ensure timely MoF declarations. Document business purpose. Review OHADA transfer pricing rules with local counsel.',
     'Corp Treasury / Legal',
     'Monitoring',
     'PPT Slide 22'),
]

# Summary counts
sev_counts = {}
cat_counts = {}
status_counts = {}
for b in blockers:
    sev_counts[b[3]] = sev_counts.get(b[3], 0) + 1
    cat_counts[b[2]] = cat_counts.get(b[2], 0) + 1
    status_counts[b[7]] = status_counts.get(b[7], 0) + 1

# Summary paragraph
p = doc.add_paragraph()
p.add_run(f'Total Blockers: {len(blockers)}').bold = True
p.add_run(f'  |  Critical: {sev_counts.get("CRITICAL",0)}  |  High: {sev_counts.get("HIGH",0)}  |  Medium: {sev_counts.get("MEDIUM",0)}  |  Low: {sev_counts.get("LOW",0)}')
doc.add_paragraph()

# Helper to set cell shading
def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

# Helper to set cell text with formatting
def set_cell(cell, text, bold=False, size=Pt(9), color=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.size = size
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color

# Create table
headers = ['ID', 'Blocker', 'Category', 'Severity', 'Impact', 'Mitigation / Action', 'Owner', 'Status', 'Source']
col_widths = [Cm(1.5), Cm(3.5), Cm(2.0), Cm(1.8), Cm(5.5), Cm(5.5), Cm(2.8), Cm(1.8), Cm(2.3)]

table = doc.add_table(rows=1, cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Header row
hdr = table.rows[0]
for i, h in enumerate(headers):
    set_cell(hdr.cells[i], h, bold=True, size=Pt(8), color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_shading(hdr.cells[i], '1a365d')

# Data rows
for b in blockers:
    row = table.add_row()
    cells = row.cells

    set_cell(cells[0], b[0], bold=True, size=Pt(8))
    set_cell(cells[1], b[1], bold=True, size=Pt(8))
    set_cell(cells[2], b[2], size=Pt(8), color=CAT_COLORS.get(b[2]))
    set_cell(cells[3], b[3], bold=True, size=Pt(8), color=SEV_COLORS.get(b[3]))
    set_cell(cells[4], b[4], size=Pt(8))
    set_cell(cells[5], b[5], size=Pt(8))
    set_cell(cells[6], b[6], size=Pt(8))
    set_cell(cells[7], b[7], size=Pt(8), color=STATUS_MAP.get(b[7]))
    set_cell(cells[8], b[8], size=Pt(7), color=RGBColor(0x71, 0x80, 0x96))

    # Severity row shading
    if b[3] == 'CRITICAL':
        for c in cells:
            set_cell_shading(c, 'FFF5F5')

# Set column widths
for i, w in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = w

# Legend
doc.add_paragraph()
doc.add_heading('Severity Definitions', level=2)
legend = doc.add_table(rows=5, cols=2)
legend.style = 'Table Grid'
set_cell(legend.rows[0].cells[0], 'Severity', bold=True, size=Pt(9))
set_cell(legend.rows[0].cells[1], 'Definition', bold=True, size=Pt(9))
set_cell_shading(legend.rows[0].cells[0], '1a365d')
set_cell_shading(legend.rows[0].cells[1], '1a365d')
legend.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
legend.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

defs = [
    ('CRITICAL', 'Blocks implementation or creates unacceptable risk. Must be resolved before proceeding to target state design.'),
    ('HIGH', 'Significant constraint that limits design options or creates material delay risk. Requires active mitigation.'),
    ('MEDIUM', 'Notable constraint that needs to be managed but does not block progress. Monitor and plan around.'),
    ('LOW', 'Minor issue or legacy item. Address opportunistically.'),
]
for i, (sev, defn) in enumerate(defs, 1):
    set_cell(legend.rows[i].cells[0], sev, bold=True, size=Pt(9), color=SEV_COLORS[sev])
    set_cell(legend.rows[i].cells[1], defn, size=Pt(9))

# Save
out = 'C:/Users/adria/ctm-transactions/GMA Blocker Register.docx'
doc.save(out)
print(f'Saved: {out}')
